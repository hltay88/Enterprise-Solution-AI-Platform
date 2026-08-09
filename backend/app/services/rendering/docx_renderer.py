"""DOCX rendering for deliverables (ATLAS-044)."""

from __future__ import annotations

from io import BytesIO
from typing import Any


def render_document_docx(
    *,
    title: str,
    status: str,
    sections: list[dict[str, Any]],
    document_label: str = "Document",
) -> bytes:
    from docx import Document

    doc = Document()
    if str(status).lower() != "approved":
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "DRAFT — NOT APPROVED FOR CUSTOMER RELEASE"
    doc.add_heading(title or document_label, level=0)
    for section in sections:
        doc.add_heading(str(section.get("title") or section.get("section_type")), level=1)
        for item in section.get("content_items") or []:
            text = str(item.get("text") or "").strip()
            if not text:
                continue
            content_type = str(item.get("content_type") or "paragraph")
            if content_type == "speaker_notes":
                continue
            prefix = ""
            if item.get("review_required"):
                prefix = "[REVIEW REQUIRED] "
            if content_type == "bullet_list":
                doc.add_paragraph(prefix + text, style="List Bullet")
            else:
                doc.add_paragraph(prefix + text)
        assumptions = section.get("assumptions") or []
        if assumptions:
            doc.add_paragraph("Assumptions:", style="Intense Quote")
            for assumption in assumptions:
                doc.add_paragraph(str(assumption), style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_proposal_docx(
    *,
    title: str,
    status: str,
    sections: list[dict[str, Any]],
) -> bytes:
    """Backward-compatible alias for document DOCX rendering."""
    return render_document_docx(
        title=title,
        status=status,
        sections=sections,
        document_label="Proposal",
    )
